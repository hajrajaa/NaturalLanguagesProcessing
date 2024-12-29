#%%
import os
import re
#import pandas as pd 
from docx import Document
import sys

import json


#%%

numbers={
    'אחת':1,
    'שתיים':2,
    'שלוש':3,
    'ארבע':4,
    'חמש':5,
    'שש':6,
    'שבע':7,
    'שמונה':8,
    'תשע':9,
    'עשר':10,

    'עשרים':20,
    'שלושים':30,
    'ארבעים':40,
    'חמישים':50,
    'שישים':60,
    'שבעים':70,
    'שמונים':80,
    'תשעים':90,

    'מאה':100,
    'מאתיים':200

}


postions_fields={ 

    'הכנסת','הכלכלה','התכנון','החינוך','התרבות','התעשייה','המסחר','האוצר','איכות הסביבה','הספורט','המשפטים','העבודה',
    'הרווחה','התחבורה','המשטרה','הבריאות','החקלאות',
    'התקשורת','ראש הממשלה','הפנים','קליטת העלייה','התיירות','ענייני דתות','התעסוקה','ביטחון פנים','התשתיות הלאומיות',
    'פיתוח הכפר','הבינוי','השיכון','המדע','הטכנולוגיה','ביטחון','החוץ','הבטיחות בדרכים',
    'הגנת הסביבה','קליטת העלייה','נושאים אסטרטגיים','ענייני מודיעין','אזרחים ותיקים','המודיעין','האנרגיה','המים','העלייה','הקליטה','השירותים החברתיים',
    'הביטחון','הפרלמנט האירופי','התשתיות','פנים','הלאומיות','ועדת'

    }



postions_keywords={
    'סגן','סגנית','שר','שרת','מזכיר','מזכירת','השר','השרה','המשנה','יהיו"ר','היו"ר','תשובת','אורח','אורחת','דובר','דוברת','יור','יו"ר','מר'
    'מ"מ','היו”ר','במשרד','ד"ר','עו"ד','נצ"מ','ניצב','שופט','מל','נשיא','מ"מ','רשף','טפסר משנה','מר','דובר_המשך',"פרופ'","'פרופ"
    ,'1היו"ר','ייור','היו"רי','<< יור >>'


}

skip={

    'קריאה','קריאות','נכחו','סדר היום','חברי הוועדה','חברי','מוזמנים','ייעוץ משפטי','מנהלת הוועדה','רישום פרלמנטרי',
    'משתתפים','נושא','רישום','מנהל/ת הוועדה','מנהל הוועדה','דיון','יועצים','רכזת'
    ,'כותבת','הצעת','מנהלי הוועדה','הוועדה','נרשם ע"י','הצעת','הספרות','רשמת','רצח','החלטת','החלטה','יועצת','הישבה ננעלה'
    'פרלמנטארית','הצעות','הישיבה','יום','הטקס','יועץ','קצרנית','מנחה','נוכחים','ברכות','הרצאה','סדר-היום','רשמה','הצגת','פתיחה','סדרן'

}   




# to convert hebrew numbers to numbers
def from_hebrew_to_number(text):

    # split the text by 's' or spaces
    parts=re.split(r'[-\s]',text)

    sum=0
    prev_num=0

    for part in parts:
        
        if part.startswith(('ו','ה')):
            part=part[1:]

        if part=='מאות':
            sum-=prev_num
            sum+=prev_num*100
            prev_num=0
            continue

        if part=='עשרה':
            sum+=10
            prev_num=0
            continue

        if part in numbers:
            prev_num=numbers[part]
            sum+=prev_num
            
    return sum 

# to check if the pargraph is underlined 
def is_underlined(paragraph):

    par_style=paragraph.style

    while par_style:
        if par_style and par_style.font.underline:
            return True

        par_style=par_style.base_style
    for run in paragraph.runs:
        if run.underline:
            return True
        
    # no underline found  
    return False


def is_bold(paragraph):
    
    par_style=paragraph.style

    while par_style:
        if par_style and par_style.font.bold:
            return True

        par_style=par_style.base_style
    for run in paragraph.runs:
        if run.bold:
            return True
        
    # not bold paragraph 
    return False

    



#%% classes 

class Sentence:

    def __init__(self,speaker_name,sentence_text):
        self.speaker_name=speaker_name
        self.sentence_text=sentence_text
        self.toknes=self.get_toknes()
        self.length=self.toknes_length()


    def get_toknes(self):

        punctuation={'.',',','?','!',';',':','"','(',')','-','%'}
        exluded_punc={':','\\','/','-'}
        toknes=""

        curr_tokne=""
        length=len(self.sentence_text)

        for i,char in enumerate(self.sentence_text):

            # handle special cases 
            #(7,500)
            if char in [',','.'] and i>0 and i+1 < length and self.sentence_text[i-1].isdigit() and self.sentence_text[i+1].isdigit():
                curr_tokne+=char
                continue

            # case : 7:00
            if char.isdigit() and i+1< length and self.sentence_text[i+1] in exluded_punc:
                curr_tokne+=char 
                continue
            # case: 05.
            elif char in exluded_punc and curr_tokne and curr_tokne[-1].isdigit():
                curr_tokne+=char
                continue
            #.05
            elif char.isdigit() and curr_tokne and curr_tokne[-1] in exluded_punc:
                curr_tokne+=char
                continue
            # case : מ"מ
            elif char in['”','"','\''] and i>0 and i+1<length and (
                '\u0590'<=self.sentence_text[i-1]<='\u05EA') and ('\u0590'<=self.sentence_text[i+1]<='\u05EA'):
                curr_tokne+=char
                continue
           


            if char in punctuation:

                if curr_tokne:
                    toknes+=curr_tokne +" "
                    curr_tokne=""


                # set the punctuation as a individual tokne 
                toknes+=char +" "
                continue

            if char.isspace():
                
                if curr_tokne:
                    toknes+=curr_tokne +" "
                    curr_tokne=""
                
            else:
                    curr_tokne+= char 

        ## add the last toknee 
        if curr_tokne:
            toknes+=curr_tokne +" "
            #curr_tokne=""

        return toknes.strip() 
        
    def toknes_length(self):

        toknes=self.toknes.split()
        return len(toknes)
    
    
    
        

    

class Protocol:

    def __init__(self,file_name,file_path):


        data = Protocol.extract_protocol_data(file_name)
        

        self.knesset_num,self.protocol_type,self.file_name=data
    
        

        self.protocol_num=Protocol.extract_protocol_num(file_path)

        self.speakers_text=[]

        #self.speakers_text=self.extract_sentences(file_path)

        #self.all_sentences=self.get_all_sentences()


        #self.sentences=self.extract_sentences(file_path)


        # #add a list of Sentence objects to the Protocol object 
        # self.sentences = Sentence(file_path)

        #self.speakers=Protocol.get_speakers_names(file_path)

    #task 1.1
    #Extract data from protocol file names
    @staticmethod
    def extract_protocol_data(file):

        valid_format=re.match(r'(\d{2})_(ptm|ptv)_(.+).docx',file)

        if valid_format:
            
                knesset_num = int(valid_format.group(1))
                protocol_type= "plenary" if valid_format.group(2)=="ptm" else "committee"
                file_name=valid_format.group(3)
                return knesset_num,protocol_type,file_name
        else:
                
                return None
    
    
    # task 1.2
    @staticmethod 
    def extract_protocol_num(file_path):

        try:

            curr_doc=Document(file_path)
            

            patterns=[
                r"פרוטוקול מס['\"]?\s*(\d+)",
                r"ישיבה מס['\"]?\s*(\d+)",
                r"הישיבה\s+([\w-]+)\s+של"
            ]

            for par in curr_doc.paragraphs:

                for pattern in patterns:
                    
                    match=re.search(pattern,par.text)
                    
                    if match:
                        
                        if pattern==patterns[2]:
                            return from_hebrew_to_number(match.group(1))
                        else:
                            return int(match.group(1))
                        
            # no match found            
            return -1 
        
        except Exception as e:
            return -1 
        

    # task 1.3 
    @staticmethod
    def extract_sentences(file_path):

        curr_doc=Document(file_path)

        speakers_text=[]


        for par_index,par in enumerate (curr_doc.paragraphs):
            
            if is_underlined(par) :
                
                
                # remove white spaces
                match=re.match(r"^(.*?)\:",par.text.strip())

                if match :

                    
                    new_speaker=match.group(1).strip()
                    text=Protocol.extract_speaker_text(par_index+1,curr_doc)
                    new_speaker=Protocol.filter_speakrs_names(new_speaker)

                    
                    if new_speaker is not None:


                        sentences=Protocol.divide_text_into_sentences(text)
                       

                        for sentence in sentences:
                            if Protocol.filter_sentence(sentence) :

                                new_sentence=Sentence(new_speaker,sentence)
                                if new_sentence.length>4:
                                    speakers_text.append(new_sentence)          
                  
        return speakers_text
    
    @staticmethod
    def extract_speaker_text(par_index,documnet):

        # speaker_text=[] 
        # paragraphs=documnet.paragraphs

        # for par in paragraphs[par_index:]:
            
        #     if par.text.strip() and (
        #     #par.text.startswith("הישיבה ננעלה") or 
        #     par.text.startswith("הצבעה מס'") or 
        #     par.text.startswith("הצבעה") or 
        #     par.text.startswith("קריאה") or
        #     is_underlined(par) or 
        #     is_bold(par)):
        #         break

        #     if par.text.strip():
        #         speaker_text.append(par.text.strip())

        # return "".join(speaker_text)

        speaker_text="" 

        for par in documnet.paragraphs[par_index:]:
            

            if par.text.startswith("\"") and par.text.endswith("\""):
                break
            #if (par.text.startswith("\"") and par.text.endswith("\"")) or is_underlined(par) or is_bold(par):
            if is_underlined(par) or is_bold(par):
                if par.text.strip():
                    break
            elif par.text.startswith("הישיבה ננעלה"): 
                break
            elif par.text.startswith("הצבעה מס'"):
                continue 
            speaker_text+= par.text.strip() 

        return speaker_text
    
    @staticmethod
    def filter_speakrs_names(name):
        
        fileds=[]
        sorted_keywords= sorted(postions_keywords,key=len,reverse=True)

        filtered_name=re.sub(r'\(.*?\)','',name).strip()


        ## check later if it is better to include it in the list instead 

        filtered_name= re.sub(r'<<.*?>>|<<.*?<<|>>.*?>>','',filtered_name.strip())


        
        if any(filtered_name.startswith(word)for word in skip) or any( word in filtered_name for word in skip )  :
             return None   

        
        while filtered_name.startswith(('>','<')):
            filtered_name=filtered_name[1:]
            if not filtered_name:
                return None

     
        for keyword in sorted_keywords:

            pattern=rf'(?:^|\s){re.escape(keyword)}(?:\s|$)'
            if re.search(pattern,filtered_name):
               filtered_name=re.sub(pattern,' ',filtered_name).strip()
               


        for field in postions_fields:
            field=re.escape(field)
            fileds.append(field)

    

        multy_feilds='|'.join(f"(?:ו?ל?)?{field}(?:ול|,|ל|ו)?" for field in fileds)

        fields_pattern=r'\b(?:'+ multy_feilds +r')\b'
        
        
        filtered_name=re.sub(fields_pattern,'',filtered_name)
        filtered_name=re.sub(r'\s*,\s*','',filtered_name)


            
        # remove extra white spaces
        filtered_name=re.sub(r'\s{2,}',' ',filtered_name)


        return filtered_name.strip()
    
    # # task 1.4 
    # def get_all_sentences(self):

    #     all_sentences=[]

    #     for speaker,text in self.speakers_text.items():

    #         sentences=Protocol.divide_text_into_sentences(text)

    #         for sentence in sentences:

    #             all_sentences.append(Sentence(speaker,sentence))

    #     return all_sentences 

    # def get_all_speakers(self):

    #     speakers=[]

    #     for sentence in self.speakers_text:
    #         if sentence.speaker_name not in speakers:
    #             speakers.append(sentence.speaker_name)

    #     return speakers 

    @staticmethod                 
    def divide_text_into_sentences(text):

         end_signs=['.','?','!','".','--','---']

         prev_sent=""

         sentences=[]

         # remove extra white spaces
         text=re.sub(r'\s{2,}','',text)

         i=0
         while i< len(text):
             char=text[i]
             prev_sent+= char 
             
             if any (prev_sent.endswith(sign) for sign in end_signs):
                 
                 if len(prev_sent)>1 and prev_sent[-2].isdigit() and (i+1 <len(text) and ((text[i+1].isdigit()) or (text[i+1] in "".join(end_signs)))):
                        i+=1
                        continue
                 if 1<len(prev_sent)<=2 and char=='.' and  '\u0590'<=prev_sent[-2]<='\u05EA':
                        i+=1
                        continue
                 if Protocol.filter_sentence(prev_sent):
                     sentences.append(prev_sent)
                 prev_sent=""
             i+=1
         if prev_sent:
             if Protocol.filter_sentence(prev_sent):
                     sentences.append(prev_sent)
         return sentences
    

    # task 1.5
    @staticmethod 
    def filter_sentence(sentence):

        invalid_signs=['---','- - -','--','- -','...','. . .','__','___','_ _','_ _ _','<<','>>']
        
        # short sentence 
        if len(sentence)<=2:
            return False
        # invalid signs 
        elif any(sign in sentence for sign in  invalid_signs):
            return False
        
        # check if there is no hebrew charcters in the sentence 
        elif not re.search(r'[א-ת]',sentence):
            return False 
                                                             
        
        # word in english 
        elif  re.search(r'[a-zA-Z]',sentence):
            return False 
        
    
        return True 
    

    # # task 1.7
    # def extract_valid_sentences(self):

    #     for sentence in self.all_sentences:
    #         if not Protocol.filter_sentence(sentence.sentence_text):
    #             self.all_sentences.remove(sentence)

    #         elif len(sentence.toknes) >= 4 :
    #             self.all_sentences.remove(sentence)

        





def extract_corpus(file_path):

    corpus=[]
    for file_name in os.listdir(file_path):

        protocol=Protocol(file_name,file_path)
        sentences=protocol.extract_sentences(file_path)

        for sentence in sentences:
            if sentence is not None:
                line={
                        "protocol_name":protocol.file_name,
                    "knesset_number":protocol.knesset_num,
                    "protocol_type":protocol.protocol_type,
                    "protocol_number":protocol.protocol_num,
                    "spekaer_name":sentence.speaker_name,
                    "sentence_text":sentence.sentence_text
                }
                corpus.append(line)
    return corpus




#%% Main code 

if __name__ == "__main__":

    try:

        if len(sys.argv)==3:

            file_path=sys.argv[1]
            output_file=sys.argv[2]

            corpus=extract_corpus(file_path)

            with open(output_file,"w",encoding="utf-8") as file:
                for line in corpus:
                    line.write(json.dumps(line,ensure_ascii=False)+"\n")

    except Exception as e:
        raise e
        



#%% Main code

if __name__ == "__main__":
    
    path=r'knesset_protocols'

    
    print(path)
    if not(os.path.exists(path)):
        print("The path does not exist")


        
    print(os.listdir(path))
    output_file="knesset_corpus.jsonl"

    with open (output_file,"w",encoding="utf-8") as jsonl_file:

        for file_name in os.listdir(path):

            if file_name.endswith(".docx"):

                file_path=os.path.join(path,file_name)
                protocol=Protocol(file_name,file_path)
                sentences=protocol.extract_sentences(file_path)


                for sentence in sentences:

                    json_line={
                        "protocol_name":protocol.file_name,
                        "knesset_number":protocol.knesset_num,
                        "protocol_type":protocol.protocol_type,
                        "protocol_number":protocol.protocol_num,
                        "speaker_name":sentence.speaker_name,
                        "sentence_text":sentence.sentence_text
                    }
                    jsonl_file.write(json.dumps(json_line,ensure_ascii=False)+"\n")


 



# %%
