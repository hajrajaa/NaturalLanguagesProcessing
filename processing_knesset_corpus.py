#%%
import zipfile
import os
import re
import pandas as pd 
from docx import Document

#import json


#%%

numbers={
    'אחת':1,
    'שתיים':2,
    'שלוש':3,
    'ארבע':4,
    'חמש':5,
    'שש':6,
    'שבע':7,
    'שמונה':8,
    'תשע':9,
    'עשר':10,

    'עשרים':20,
    'שלושים':30,
    'ארבעים':40,
    'חמישים':50,
    'שישים':60,
    'שבעים':70,
    'שמונים':80,
    'תשעים':90,

    'מאה':100,
    'מאתיים':200

}


postions_fields={ 

    'הכנסת','הכלכלה','התכנון','החינוך','התרבות','התעשייה','המסחר','האוצר','איכות הסביבה','הספורט','המשפטים','העבודה',
    'הרווחה','התחבורה','המשטרה','הבריאות','החקלאות',
    'התקשורת','ראש הממשלה','הפנים','קליטת העלייה','התיירות','ענייני דתות','התעסוקה','ביטחון פנים','התשתיות הלאומיות',
    'פיתוח הכפר','הבינוי','השיכון','המדע','הטכנולוגיה','ביטחון','החוץ','הבטיחות בדרכים',
    'הגנת הסביבה','קליטת העלייה','נושאים אסטרטגיים','ענייני מודיעין','אזרחים ותיקים','המודיעין','האנרגיה','המים','העלייה','הקליטה','השירותים החברתיים',
    'הביטחון','הפרלמנט האירופי','התשתיות','פנים','הלאומיות',''
    # ,
    # 'והספורט','הספורט','והתרבות ',''
    # '','','','','','','','','','','','','','','','','',''
    }



postions_keywords={
    'סגן','סגנית','שר','שרת','מזכיר','מזכירת','השר','השרה','המשנה','היו"ר','תשובת','אורח','אורחת','דובר','דוברת','יור','פרופ','יו"ר','מר'
    'מ"מ','היו”ר','במשרד','ד"ר','עו"ד','נצ"מ','ניצב','שופט','מל','נשיא','מ"מ','רשף','טפסר משנה','מר','פרופ\''



}

skip={

    'קריאה','קריאות','נכחו','סדר היום','חברי הוועדה','חברי','מוזמנים','ייעוץ משפטי','מנהלת הוועדה','רישום פרלמנטרי',
    'משתתפים','נושא','רישום','מנהל/ת הוועדה','מנהל הוועדה','דיון','יועצים','רכזת'
    ,'כותבת','הצעת','מנהלי הוועדה','הוועדה','נרשם ע"י','הצעת','הספרות','רשמת','רצח','החלטת','החלטה','יועצת','הישבה ננעלה'
    'פרלמנטארית','הצעות','הישיבה','יום','הטקס','יועץ','קצרנית','מנחה','נוכחים','ברכות','הרצאה','סדר-היום','רשמה','הצגת'

}   




# to convert hebrew numbers to numbers
def from_hebrew_to_number(text):

    # split the text by 's' or spaces
    parts=re.split(r'[-\s]',text)

    sum=0
    prev_num=0

    for part in parts:
        
        if part.startswith(('ו','ה')):
            part=part[1:]

        if part=='מאות':
            sum-=prev_num
            sum+=prev_num*100
            prev_num=0
            continue

        if part=='עשרה':
            sum+=10
            prev_num=0
            continue

        if part in numbers:
            prev_num=numbers[part]
            sum+=prev_num
            
    return sum 

# to check if the pargraph is underlined 
def is_underlined(paragraph):

    par_style=paragraph.style

    while par_style:
        if par_style and par_style.font.underline:
            return True

        par_style=par_style.base_style
    for run in paragraph.runs:
        if run.underline:
            return True
        
    # no underline found  
    return False



#%% classes 

class Sentence:

    def __init__(self,speaker_name,sentence_text):
        self.speaker_name=speaker_name
        self.sentence_text=sentence_text
        self.length=len(sentence_text)
        self.toknes=self.get_toknes()

    def get_toknes(self):

        punctuation={'.',',','?','!',';',':','"',"'",'(',')','-'}
        exluded_punc={':','\\','/','-'}
        toknes=[]

        curr_tokne=""
        length=len(self.sentence_text)

        for i,char in enumerate(self.sentence_text):

            # hancle special cases 
            #(7,500)
            if char=="," and i>0 and i+1 < length and self.sentence_text[i-1].isdigit() and self.sentence_text[i+1].isdigit():
                curr_tokne+=char
                continue

            # case : 7:00
            if char.isdigit() and i+1< length and self.sentence_text[i+1] in exluded_punc:
                curr_tokne+=char 
                continue
            # case: 05.
            elif char in exluded_punc and curr_tokne and curr_tokne[-1].isdigit():
                curr_tokne+=char
                continue
            #.05
            elif char.isdigit() and curr_tokne and curr_tokne[-1] in exluded_punc:
                curr_tokne+=char
                continue
            # case : מ"מ
            elif char in['”','"'] and i>0 and i+1<length and (
                '\u0590'<=self.sentence_text[i-1]<='\u05EA') and ('\u0590'<=self.sentence_text[i+1]<='\u05EA'):
                curr_tokne+=char
                continue


            if char in punctuation:

                if curr_tokne:
                    toknes.append(curr_tokne)
                    curr_tokne=""


                # set the punctuation as a individual tokne 
                toknes.append(char)

            elif char.isspace():
                
                if curr_tokne:
                    toknes.append(curr_tokne)
                    curr_tokne=""
                #toknes.append(" ")
            else:
                    curr_tokne+= char 

        ## add the last toknee 
        if curr_tokne:
            toknes.append(curr_tokne)

        return toknes 
    


        

            



        

    

class Protocol:

    def __init__(self,file_name,file_path):


        data = Protocol.extract_protocol_data(file_name)
        

        self.knesset_num,self.protocol_type,self.file_name=data
    
        

        self.protocol_num=Protocol.extract_protocol_num(file_path)

        self.speakers_text=self.extract_sentences(file_path)

        self.all_sentences=self.get_all_sentences()


        #self.sentences=self.extract_sentences(file_path)


        # #add a list of Sentence objects to the Protocol object 
        # self.sentences = Sentence(file_path)

        #self.speakers=Protocol.get_speakers_names(file_path)

    #task 1.1
    #Extract data from protocol file names
    @staticmethod
    def extract_protocol_data(file):

        valid_format=re.match(r'(\d{2})_(ptm|ptv)_(.+).docx',file)

        if valid_format:
            
                knesset_num = int(valid_format.group(1))
                protocol_type= "plenary" if valid_format.group(2)=="ptm" else "committee"
                file_name=valid_format.group(3)
                return knesset_num,protocol_type,file_name
        else:
                
                return None
    
    
    
    # task 1.2
    @staticmethod 
    def extract_protocol_num(file_path):

        try:

            curr_doc=Document(file_path)
            

            patterns=[
                r"פרוטוקול מס['\"]?\s*(\d+)",
                r"ישיבה מס['\"]?\s*(\d+)",
                r"הישיבה\s+([\w-]+)\s+של"
            ]

            for par in curr_doc.paragraphs:

                for pattern in patterns:
                    
                    match=re.search(pattern,par.text)
                    
                    if match:
                        
                        if pattern==patterns[2]:
                            return from_hebrew_to_number(match.group(1))
                        else:
                            return int(match.group(1))
                        
            # no match found            
            return -1 
        
        except Exception as e:
            return -1 
        

    # task 1.3 
    def extract_sentences(self,file_path):

        curr_doc=Document(file_path)

        speakers_text={}


        for par_index,par in enumerate (curr_doc.paragraphs):

            if is_underlined(par):
                
                # remove white spaces
                match=re.match(r"^(.*?)\:",par.text.strip())

                if match :
                    
                    new_speaker=match.group(1).strip()

                    if any( word in new_speaker for word in skip ) or any(new_speaker.startswith(word)for word in skip) :
                        continue  

                    text=self.extract_speaker_text(par_index+1,curr_doc)
                    #print(text)
                    new_speaker=self.filter_speakrs_names(new_speaker)

                    if new_speaker in speakers_text:
                        speakers_text[new_speaker]+="\n"+text
                    else:
                        speakers_text[new_speaker]=text
                   
                  
        return speakers_text
    
    @staticmethod
    def extract_speaker_text(par_index,documnet):

        speaker_text="" 

        for par in documnet.paragraphs[par_index:]:

            if is_underlined(par):
                break
            speaker_text+= par.text.strip() +"\n"

        return speaker_text
    
    @staticmethod
    def filter_speakrs_names(name):


        keywords=[]
        fileds=[]

        
        filtered_name=re.sub(r'\(.*?\)','',name).strip()


        ## check later if it is better to include it in the list instead 

        filtered_name= re.sub(r'<<.*?>>|<<.*?<<|>>.*?>>','',filtered_name)

        
        while filtered_name.startswith(('>','<')):
            filtered_name=filtered_name[1:]



        for keyword in postions_keywords:

            keyword=re.escape(keyword)
            keywords.append(keyword)

        for field in postions_fields:
            field=re.escape(field)
            fileds.append(field)

        
        keyword_pattern=r'\b(?:' +'|'.join(keywords)+r')\b'
        keyword_pattern=fr'{keyword_pattern}(?:\s*ל|\b[\s\-]*)?'
        


        multy_feilds='|'.join(f"(?:ו?ל?)?{field}(?:ול|,|ל|ו)?" for field in fileds)

        fields_pattern=r'\b(?:'+ multy_feilds +r')\b'

        prev_filtered_name=None
        while filtered_name!=prev_filtered_name:
            prev_filtered_name=filtered_name
            filtered_name=re.sub(keyword_pattern,'',filtered_name)
            filtered_name=re.sub(fields_pattern,'',filtered_name)
            filtered_name=re.sub(r'\s*,\s*','',filtered_name)
            

            
        # remove extra white spaces
        filtered_name=re.sub(r'\s{2,}','',filtered_name)


        return filtered_name.strip()
    
    # task 1.4 
    def get_all_sentences(self):

        all_sentences=[]

        for speaker,text in self.speakers_text.items():

            sentences=Protocol.divide_text_into_sentences(text)

            for sentence in sentences:

                all_sentences.append(Sentence(speaker,sentence))

        return all_sentences 

                       
    @staticmethod                 
    def divide_text_into_sentences(text):

        beg_signs=['א','ב','ג','ד','ה','ו','ז','ח','ט','י','כ']

        end_signs=['.','?','!','".','--','---']

        prev_sent=""

        sentences=[]

        # remove extra white spaces
        text=re.sub(r'\s{2,}','',text)

        for char in text:
            prev_sent+=char

            #prev_sent=prev_sent.strip()

            if char in end_signs:

                
                if  prev_sent[-2].isdigit():
                    continue
                elif char=='.' and prev_sent[-2] in beg_signs:
                    continue
                elif Protocol.filter_sentence(prev_sent):
                    sentences.append(prev_sent)
                    prev_sent=""

        if prev_sent and Protocol.filter_sentence(prev_sent):
              sentences.append(prev_sent)

        return sentences
    
    # task 1.5
    @staticmethod 
    def filter_sentence(sentence):

        invalid_signs=['---','- - -','--','- -','...','. . .','__','___','_ _','_ _ _']
        
        # short sentence 
        if len(sentence)<=10:
            return False
        # invalid signs 
        elif any(sign in sentence for sign in  invalid_signs):
            return False
        
        # word in english 
        elif re.match(r"^[a-zA-Z\-']+$",sentence):
            return False 
        
        # Toooo Doooooo# 
        # check if there is no hebrew charcters in the sentence 
        elif not re.search(r'[א-ת]',sentence):
            return False 
    
        return True 
    

    # task 1.6 


        





            



        








#%% Main code

if __name__ == "__main__":
    
    path=r"Knesset_protocols\protocol_for_hw1"


    # file_name="17_ptm_533398.docx"

    # protocol=Protocol(file_name,os.path.join(path,file_name))
    # for sentence in protocol.all_sentences:
    #     print("speaker:",sentence.speaker_name)
    #     print("text:",sentence.sentence_text)

    text="לחבר הכנסת - הפתרון הנקודתי שנתנו לפתחת-ניצנה,"
    text1="המחיר הוא 6,500 ש\"ח, התאריך: 01/01/2024, והשעה 10:50."
    s=Sentence("לולו",text1)

    toknes=Sentence.get_toknes(s)
    print(toknes)


    # text="ממנה?\""

    # # text="א. אני רגאגאאגאג"
    # for sentence in Protocol.divide_text_into_sentences(text):
    #     print(sentence)
    #     print("\n")

    #print(Protocol.divide_text_into_sentences(text))


    # for file_name in os.listdir(path):
    #     if file_name.endswith('.docx'):
    #         print(file_name)
    #         protocol=Protocol(file_name,os.path.join(path,file_name))
    #         #print(protocol.knesset_num,protocol.protocol_type,protocol.file_name)
            
    #         for sentence in protocol.sentences:
    #             if sentence in None:
    #                 print("looooooooooooo")
    #             print("speaker:",sentence.speaker)
    #             print("text:",sentence.text)
                

    # speaker='סגן שרת החינוך והתרבות וליד גולדמן'
    # speaker1=' סגן שר החינוך , התרבות והספורט'

    # print(filter_speakrs_names(speaker))
    # print(filter_speakrs_names(speaker1))

    # speaker=['השר לקליטת העלייה','שר הכלכלה','סגן השר לביטחון','סגנית מזכיר הכנסת']

    # filtered_speaker=filter_speakrs_names(speaker)
    # print(filtered_speaker)


    # for file_name in os.listdir(path):
    #      if file_name.endswith('.docx'):        
    #          print(extract_protocol_num(os.path.join(path,file_name)))
    
    

  


    # for file_name in os.listdir(protocol_files):

    #     protocol=extract_protocol_data(file_name)
    #     if protocol:
    #         print(protocol.knesset_num,protocol.protocol_type,protocol.file_name)
    #     else:
    #         print(f"invalid file name: {file_name}")




# %%
